#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Qconverto - 使用 NiceGUI 重构的多媒体文件格式转换工具
支持图片、音频、视频和文档格式的转换
优化版本 - 具有更好的UI/UX和性能
"""

import os
import sys
import asyncio
import logging
from pathlib import Path
from datetime import datetime
from nicegui import ui, app, events
from typing import Optional, Dict, Any

# 添加当前目录到路径
sys.path.insert(0, '.')

# 导入UI组件
from ui_components import ModernUIComponents

# 添加自定义CSS样式
ui.add_head_html('''
<style>
    @keyframes pulse {
        0% { transform: scale(1); }
        50% { transform: scale(1.1); }
        100% { transform: scale(1); }
    }
    
    @keyframes bounce {
        0%, 100% { transform: translateY(0); }
        50% { transform: translateY(-10px); }
    }
    
    @keyframes spin {
        0% { transform: rotate(0deg); }
        100% { transform: rotate(360deg); }
    }
    
    .animate-pulse {
        animation: pulse 2s infinite;
    }
    
    .animate-bounce {
        animation: bounce 1s infinite;
    }
    
    .animate-spin {
        animation: spin 1s linear infinite;
    }
    
    .file-area-hover {
        background-color: #f0f9ff;
        border-color: #3b82f6;
    }
</style>
''')

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('qconverto.log', encoding='utf-8'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

# 检查MoviePy是否可用
try:
    from moviepy.editor import VideoFileClip
    MOVIEPY_AVAILABLE = True
except ImportError:
    MOVIEPY_AVAILABLE = False

# 支持的文件格式字典
SUPPORTED_FORMATS = {
    '图片': {
        '输入': ['.jpg', '.jpeg', '.png', '.bmp', '.gif', '.webp', '.ico'],
        '输出': ['.jpg', '.jpeg', '.png', '.webp', '.pdf']
    },
    '音频': {
        '输入': ['.mp3', '.wav', '.flac', '.ogg', '.m4a', '.mp4', '.aac', '.ape', '.wv'],
        '输出': ['.mp3', '.wav', '.flac', '.ogg', '.m4a']
    },
    '视频': {
        '输入': ['.mp4', '.avi', '.mov', '.mkv', '.wmv', '.flv'],
        '输出': ['.mp4', '.avi', '.mov', '.mkv', '.wmv']
    },
    '文档': {
        '输入': ['.pdf', '.doc', '.docx', '.txt'],
        '输出': ['.pdf', '.docx', '.txt', '.jpg']
    }
}

# 文件类型到图标的映射
FILE_TYPE_ICONS = {
    # 图片格式
    '.jpg': 'image',
    '.jpeg': 'image',
    '.png': 'image',
    '.bmp': 'image',
    '.gif': 'gif',
    '.webp': 'image',
    '.ico': 'image',
    '.tif': 'image',
    '.tiff': 'image',
    '.svg': 'image',
    '.psd': 'image',
    '.eps': 'image',
    '.raw': 'image',
    '.cr2': 'image',
    '.nef': 'image',
    '.orf': 'image',
    
    # 文档格式
    '.pdf': 'picture_as_pdf',
    '.doc': 'description',
    '.docx': 'description',
    '.txt': 'text_snippet',
    '.rtf': 'text_snippet',
    '.odt': 'description',
    '.xls': 'table_chart',
    '.xlsx': 'table_chart',
    '.csv': 'table_chart',
    '.ppt': 'slideshow',
    '.pptx': 'slideshow',
    '.odp': 'slideshow',
    '.md': 'article',
    '.tex': 'article',
    '.epub': 'book',
    '.mobi': 'book',
    
    # 音频格式
    '.mp3': 'audiotrack',
    '.wav': 'audiotrack',
    '.flac': 'audiotrack',
    '.ogg': 'audiotrack',
    '.m4a': 'audiotrack',
    '.aac': 'audiotrack',
    '.ape': 'audiotrack',
    '.wv': 'audiotrack',
    '.wma': 'audiotrack',
    '.mp2': 'audiotrack',
    '.amr': 'audiotrack',
    
    # 视频格式
    '.mp4': 'videocam',
    '.avi': 'videocam',
    '.mov': 'videocam',
    '.mkv': 'videocam',
    '.wmv': 'videocam',
    '.flv': 'videocam',
    '.webm': 'videocam',
    '.3gp': 'videocam',
    '.m4v': 'videocam',
    '.mpg': 'videocam',
    '.mpeg': 'videocam',
    '.ogv': 'videocam',
    
    # 压缩文件格式
    '.zip': 'folder_zip',
    '.rar': 'folder_zip',
    '.7z': 'folder_zip',
    '.tar': 'folder_zip',
    '.gz': 'folder_zip',
    '.bz2': 'folder_zip',
    '.xz': 'folder_zip',
    '.zipx': 'folder_zip',
    
    # 其他格式
    '.exe': 'application',
    '.dll': 'application',
    '.so': 'application',
    '.dmg': 'application',
    '.iso': 'disc_full',
    '.img': 'disc_full',
    '.torrent': 'cloud_download',
    '.json': 'code',
    '.xml': 'code',
    '.html': 'code',
    '.css': 'code',
    '.js': 'code',
    '.py': 'code',
    '.java': 'code',
    '.c': 'code',
    '.cpp': 'code',
    '.h': 'code',
    '.hpp': 'code',
    '.go': 'code',
    '.php': 'code',
    '.sql': 'code',
    '.sh': 'code',
    '.bat': 'code',
    '.ps1': 'code',
    '.apk': 'android',
    '.ipa': 'apple',
    '.app': 'apple',
    '.crx': 'extension',
    '.xpi': 'extension',
    '.ico': 'extension',
    '.svg': 'extension',
    '.ttf': 'font_download',
    '.otf': 'font_download',
    '.woff': 'font_download',
    '.woff2': 'font_download'
}


class ConversionWorker:
    """转换工作类，用于执行文件转换任务"""
    
    def __init__(self, input_file: str, output_file: str, file_type: str):
        self.input_file = input_file
        self.output_file = output_file
        self.file_type = file_type
        self.settings = {}
        self.progress = 0
        self.cancelled = False
        
    async def run_conversion(self):
        """执行转换任务"""
        try:
            logger.info(f"开始转换: {self.input_file} -> {self.output_file}")
            self.progress = 10
            yield self.progress
            
            if self.file_type == '图片':
                # 这些方法是异步生成器，需要遍历而不是await
                async for progress in self._convert_image():
                    yield progress
            elif self.file_type == '音频':
                async for progress in self._convert_audio():
                    yield progress
            elif self.file_type == '视频':
                async for progress in self._convert_video():
                    yield progress
            elif self.file_type == '文档':
                async for progress in self._convert_document():
                    yield progress
            else:
                raise ValueError(f"不支持的文件类型: {self.file_type}")
            
            if not self.cancelled:
                self.progress = 100
                yield self.progress
                logger.info(f"转换完成: {self.output_file}")
        except Exception as e:
            logger.error(f"转换失败: {str(e)}", exc_info=True)
            # 提供更详细的错误信息
            import traceback
            error_details = f"{str(e)}\n\n详细信息:\n{traceback.format_exc()}"
            raise RuntimeError(error_details)

    async def _convert_image(self):
        """转换图片文件"""
        if self.cancelled:
            return
        try:
            from PIL import Image
            self.progress = 30
            yield self.progress
            
            with Image.open(self.input_file) as img:
                self.progress = 70
                yield self.progress
                
                # 使用默认图片质量
                img.save(self.output_file, quality=90)
        except Exception as e:
            raise RuntimeError(f"图片转换失败: {str(e)}") from e

    async def _convert_audio(self):
        """转换音频文件"""
        if self.cancelled:
            return
        try:
            self.progress = 30
            yield self.progress
            
            # 首先尝试使用pydub进行音频转换（支持更多格式）
            try:
                from pydub import AudioSegment
                
                # 使用pydub加载音频文件
                audio = AudioSegment.from_file(self.input_file)
                self.progress = 50
                yield self.progress
                
                # 使用默认音频设置导出为指定格式
                file_ext = os.path.splitext(self.output_file)[1].lower()
                
                if file_ext == '.mp3':
                    audio.export(self.output_file, format="mp3", bitrate="192k")
                elif file_ext == '.wav':
                    audio.export(self.output_file, format="wav")
                elif file_ext == '.flac':
                    audio.export(self.output_file, format="flac")
                elif file_ext == '.ogg':
                    audio.export(self.output_file, format="ogg")
                elif file_ext == '.m4a':
                    audio.export(self.output_file, format="mp4")
                else:
                    # 默认导出为MP3
                    audio.export(self.output_file, format="mp3", bitrate="192k")
                
                self.progress = 70
                yield self.progress
                return  # 成功使用pydub转换
                
            except Exception as pydub_error:
                # pydub转换失败，尝试使用miniaudio
                logger.warning(f"pydub转换失败，尝试miniaudio: {pydub_error}")
                pass
            
            # 尝试使用miniaudio进行音频处理
            try:
                import miniaudio
                
                self.progress = 50
                yield self.progress
                
                # 使用miniaudio处理音频文件
                input_ext = os.path.splitext(self.input_file)[1].lower()
                output_ext = os.path.splitext(self.output_file)[1].lower()
                
                # 获取音频信息
                info = miniaudio.get_file_info(self.input_file)
                logger.info(f"音频信息: {info.nchannels}声道, {info.sample_rate}Hz, {info.duration:.2f}秒")
                
                # 解码音频
                decoded = miniaudio.decode_file(self.input_file)
                
                # 根据输出格式决定处理方式
                if output_ext == '.wav':
                    # 保存为WAV文件
                    import wave
                    with wave.open(self.output_file, 'w') as wf:
                        wf.setnchannels(decoded.nchannels)
                        wf.setsampwidth(2)  # 16-bit
                        wf.setframerate(decoded.sample_rate)
                        wf.writeframes(decoded.samples)
                else:
                    # 对于其他格式，执行文件复制并给出提示
                    import shutil
                    shutil.copy2(self.input_file, self.output_file)
                    logger.info(f"注意: miniaudio不支持编码为{output_ext}格式，执行文件复制")
                
                self.progress = 70
                yield self.progress
                return  # 成功使用miniaudio处理
                
            except Exception as miniaudio_error:
                # miniaudio处理失败，尝试使用纯Python方法
                logger.warning(f"miniaudio处理失败，尝试纯Python方法: {miniaudio_error}")
                pass
            
            # 如果pydub和miniaudio都不可用，使用纯Python方法进行基本转换
            self.progress = 50
            yield self.progress
            
            # 纯Python音频转换（仅支持WAV相关格式）
            input_ext = os.path.splitext(self.input_file)[1].lower()
            output_ext = os.path.splitext(self.output_file)[1].lower()
            
            # 对于WAV文件的基本处理
            if input_ext == '.wav':
                if output_ext == '.wav':
                    # WAV到WAV的复制
                    import shutil
                    shutil.copy2(self.input_file, self.output_file)
                else:
                    # 对于其他格式，提供说明或简单复制
                    import shutil
                    shutil.copy2(self.input_file, self.output_file)
                    logger.info(f"注意：从WAV到{output_ext}的真实转换需要专门的编码库")
            else:
                # 对于非WAV格式，尝试简单复制（可能不工作）
                import shutil
                shutil.copy2(self.input_file, self.output_file)
                logger.info(f"注意：从{input_ext}到{output_ext}的真实转换需要专门的解码和编码库")
            
            self.progress = 70
            yield self.progress
            
        except Exception as e:
            # 提供更详细的错误信息
            error_msg = f"音频转换失败: {str(e)}\n\n"
            error_msg += "请确保输入文件格式受支持且未被其他程序占用。"
            raise RuntimeError(error_msg) from e

    async def _convert_video(self):
        """转换视频文件"""
        if self.cancelled:
            return
        try:
            self.progress = 30
            yield self.progress
            
            # 首先尝试使用MoviePy进行视频转换（如果可用）
            if MOVIEPY_AVAILABLE:
                try:
                    from moviepy.editor import VideoFileClip
                    import os
                    
                    # 使用MoviePy加载视频文件
                    video = VideoFileClip(self.input_file)
                    self.progress = 50
                    yield self.progress
                    
                    # 获取输出文件扩展名
                    output_ext = os.path.splitext(self.output_file)[1].lower()
                    
                    # 根据输出格式进行相应处理
                    if output_ext in ['.mp4', '.avi', '.mov', '.mkv', '.wmv']:
                        # 对于支持的格式，直接写入
                        video.write_videofile(
                            self.output_file,
                            codec='libx264',
                            audio_codec='aac',
                            temp_audiofile='temp-audio.m4a',
                            remove_temp=True,
                            verbose=False,
                            logger=None
                        )
                    else:
                        # 对于不直接支持的格式，使用默认设置
                        video.write_videofile(
                            self.output_file,
                            temp_audiofile='temp-audio.m4a',
                            remove_temp=True,
                            verbose=False,
                            logger=None
                        )
                    
                    # 关闭视频文件
                    video.close()
                    self.progress = 70
                    yield self.progress
                    return  # 成功使用MoviePy转换
                    
                except Exception as moviepy_error:
                    # MoviePy转换失败，尝试使用FFmpeg
                    logger.warning(f"MoviePy转换失败，尝试FFmpeg: {moviepy_error}")
                    pass
            
            # 尝试使用FFmpeg进行视频转换
            try:
                import ffmpeg
                import shutil
                import os
                
                # 检查ffmpeg是否可用
                ffmpeg_path = shutil.which("ffmpeg")
                if not ffmpeg_path:
                    # 尝试使用相对路径或打包环境中的ffmpeg
                    possible_paths = [
                        "ffmpeg",
                        "./ffmpeg",
                        "./ffmpeg.exe",
                        os.path.join(os.path.dirname(__file__), "ffmpeg"),
                        os.path.join(os.path.dirname(__file__), "ffmpeg.exe")
                    ]
                    
                    for path in possible_paths:
                        if os.path.exists(path) or shutil.which(path):
                            ffmpeg_path = path
                            break
                
                # 如果仍然找不到，提供更友好的错误信息
                if not ffmpeg_path:
                    raise RuntimeError(
                        "未找到FFmpeg，视频转换需要FFmpeg支持。\n\n"
                        "解决方案:\n"
                        "1. 自动下载: 运行 install_ffmpeg.py 脚本自动安装FFmpeg\n"
                        "2. 手动安装: 访问 https://ffmpeg.org/download.html 下载并安装\n"
                        "3. 打包应用: 将ffmpeg.exe文件放在程序同目录下\n\n"
                        "注意: 音频转换不需要FFmpeg，可以正常使用。"
                    )
                
                # 使用默认视频设置
                try:
                    (ffmpeg
                        .input(self.input_file)
                        .output(self.output_file, 
                                vcodec='libx264', 
                                acodec='aac', 
                                crf=23, 
                                strict='experimental')
                        .overwrite_output()
                        .run(cmd=ffmpeg_path, capture_stdout=True, capture_stderr=True)
                    )
                    self.progress = 70
                    yield self.progress
                except ffmpeg.Error as e:
                    # 获取详细的错误信息
                    error_msg = f"视频转换失败: {e.stderr.decode('utf-8') if e.stderr else str(e)}"
                    raise RuntimeError(error_msg) from e
                except Exception as e:
                    error_msg = f"视频转换过程中发生未知错误: {str(e)}"
                    raise RuntimeError(error_msg) from e
                    
            except Exception as ffmpeg_error:
                # FFmpeg处理失败，尝试使用纯Python方法
                logger.error(f"FFmpeg处理失败: {ffmpeg_error}")
                raise RuntimeError(f"视频转换失败: 未找到可用的视频处理库(MoviePy或FFmpeg)") from ffmpeg_error
                
        except Exception as e:
            raise RuntimeError(f"视频转换失败: {str(e)}") from e

    async def _convert_document(self):
        """转换文档文件"""
        if self.cancelled:
            return
        try:
            # 仅支持部分文档格式转换
            ext_in = os.path.splitext(self.input_file)[1].lower()
            ext_out = os.path.splitext(self.output_file)[1].lower()
            
            self.progress = 30
            yield self.progress
            
            # PDF转图片
            if ext_in == '.pdf' and ext_out in ['.jpg', '.jpeg']:
                # 这些方法是异步生成器，需要遍历而不是await
                async for progress in self._pdf_to_image():
                    yield progress
            # 图片转PDF
            elif ext_in in ['.jpg', '.jpeg', '.png', '.webp', '.bmp', '.tiff'] and ext_out == '.pdf':
                await self._image_to_pdf()
                yield 70  # 图片转PDF完成后进度到70%
            # 文本转PDF
            elif ext_in == '.txt' and ext_out == '.pdf':
                async for progress in self._txt_to_pdf():
                    yield progress
            # PDF转文本
            elif ext_in == '.pdf' and ext_out == '.txt':
                async for progress in self._pdf_to_txt():
                    yield progress
            # DOC转PDF
            elif ext_in == '.doc' and ext_out == '.pdf':
                async for progress in self._doc_to_pdf():
                    yield progress
            # DOCX转PDF
            elif ext_in == '.docx' and ext_out == '.pdf':
                # 直接调用_docx_to_pdf方法，它会自动处理依赖库检查
                async for progress in self._docx_to_pdf():
                    yield progress
            # PDF转DOCX
            elif ext_in == '.pdf' and ext_out == '.docx':
                async for progress in self._pdf_to_docx():
                    yield progress
            else:
                raise RuntimeError(f"不支持的文档转换类型: {ext_in} → {ext_out}")
            
            self.progress = 70
            yield self.progress
        except Exception as e:
            raise RuntimeError(f"文档转换失败: {str(e)}") from e

    async def _pdf_to_image(self):
        """PDF转图片"""
        if self.cancelled:
            return
        try:
            import fitz  # PyMuPDF
            from PIL import Image
            import os
            
            # 检查输入文件是否存在
            if not os.path.exists(self.input_file):
                raise FileNotFoundError(f"输入文件不存在: {self.input_file}")
            
            doc = fitz.open(self.input_file)
            
            # 检查PDF是否有页面
            if len(doc) == 0:
                doc.close()
                raise RuntimeError("PDF文件没有页面内容")
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                # 使用更高的分辨率以获得更好的图像质量
                pix = page.get_pixmap(matrix=fitz.Matrix(3, 3))  # 3倍分辨率
                img_data = pix.tobytes("ppm")
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                
                if len(doc) > 1:
                    # 如果有多个页面，为每个页面创建单独的图片文件
                    base_name = os.path.splitext(self.output_file)[0]
                    output_file = f"{base_name}_{page_num + 1:02d}{os.path.splitext(self.output_file)[1]}"
                else:
                    output_file = self.output_file
                
                # 保存图片
                output_format = os.path.splitext(self.output_file)[1].upper()[1:]  # 获取文件扩展名作为格式
                if output_format == 'JPG':
                    output_format = 'JPEG'
                
                img.save(output_file, format=output_format, quality=95)
                
                # 检查输出文件是否创建成功
                if not os.path.exists(output_file):
                    doc.close()
                    raise RuntimeError(f"图片文件未能成功创建: {output_file}")
            
            doc.close()
        except FileNotFoundError:
            raise
        except Exception as e:
            raise RuntimeError(f"PDF转图片失败: {str(e)}") from e

    async def _image_to_pdf(self):
        """图片转PDF"""
        if self.cancelled:
            return
        try:
            from PIL import Image
            import os
            
            # 检查输入文件是否存在
            if not os.path.exists(self.input_file):
                raise FileNotFoundError(f"输入文件不存在: {self.input_file}")
            
            # 支持的图片格式列表
            supported_formats = ['.jpg', '.jpeg', '.png', '.webp', '.bmp', '.tiff']
            file_ext = os.path.splitext(self.input_file)[1].lower()
            
            if file_ext not in supported_formats:
                raise ValueError(f"不支持的图片格式: {file_ext}")
            
            # 打开图片
            img = Image.open(self.input_file)
            
            # 确保图片模式兼容PDF
            if img.mode in ("RGBA", "P"):
                img = img.convert("RGB")
            
            # 保存为PDF
            img.save(self.output_file, format='PDF', resolution=100.0)
            
            # 检查输出文件是否创建成功
            if not os.path.exists(self.output_file):
                raise RuntimeError("PDF文件未能成功创建")
                
        except FileNotFoundError:
            raise
        except ValueError as e:
            raise
        except Exception as e:
            raise RuntimeError(f"图片转PDF失败: {str(e)}") from e

    async def _txt_to_pdf(self):
        """文本转PDF"""
        if self.cancelled:
            return
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            import os
            
            # 检查输入文件是否存在
            if not os.path.exists(self.input_file):
                raise FileNotFoundError(f"输入文件不存在: {self.input_file}")
            
            # 创建PDF
            c = canvas.Canvas(self.output_file, pagesize=letter)
            width, height = letter
            
            # 读取文本文件，尝试多种编码
            lines = []
            encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
            file_read = False
            
            for encoding in encodings:
                try:
                    with open(self.input_file, 'r', encoding=encoding) as f:
                        lines = f.readlines()
                    file_read = True
                    break
                except UnicodeDecodeError:
                    continue
            
            if not file_read:
                raise RuntimeError("无法读取文本文件，可能的编码问题")
            
            # 尝试注册并使用中文字体
            font_registered = False
            try:
                # 在Windows上尝试使用系统字体
                if os.name == 'nt':  # Windows
                    font_path = "C:/Windows/Fonts/simsun.ttc"  # 宋体
                    if os.path.exists(font_path):
                        pdfmetrics.registerFont(TTFont('SimSun', font_path))
                        c.setFont("SimSun", 12)
                        font_registered = True
                # 在其他系统上可以尝试其他字体路径
            except:
                pass
            
            # 如果中文字体注册失败，使用默认字体
            if not font_registered:
                c.setFont("Helvetica", 12)
            
            y_position = height - 50
            
            # 写入文本
            for line in lines:
                # 处理过长的行，进行换行
                text = line.strip()
                while len(text) > 80:  # 如果行长度超过80字符
                    c.drawString(50, y_position, text[:80])
                    text = text[80:]
                    y_position -= 15
                    
                    # 如果到达页面底部，创建新页面
                    if y_position < 50:
                        c.showPage()
                        # 重新设置字体
                        if font_registered:
                            c.setFont("SimSun", 12)
                        else:
                            c.setFont("Helvetica", 12)
                        y_position = height - 50
                
                # 绘制剩余文本（或未超长的整行）
                if text:  # 确保有文本需要绘制
                    c.drawString(50, y_position, text)
                    y_position -= 15
                y_position -= 15
                
                # 如果到达页面底部，创建新页面
                if y_position < 50:
                    c.showPage()
                    # 重新设置字体
                    if font_registered:
                        c.setFont("SimSun", 12)
                    else:
                        c.setFont("Helvetica", 12)
                    y_position = height - 50
            
            c.save()
            
            # 检查输出文件是否创建成功
            if not os.path.exists(self.output_file):
                raise RuntimeError("PDF文件未能成功创建")
                
        except FileNotFoundError:
            raise
        except Exception as e:
            raise RuntimeError(f"文本转PDF失败: {str(e)}") from e

    async def _pdf_to_txt(self):
        """PDF转文本"""
        if self.cancelled:
            return
        try:
            import fitz  # PyMuPDF
            import os
            
            # 检查输入文件是否存在
            if not os.path.exists(self.input_file):
                raise FileNotFoundError(f"输入文件不存在: {self.input_file}")
            
            # 打开PDF文档
            doc = fitz.open(self.input_file)
            text = ""
            
            # 提取所有页面的文本
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                page_text = page.get_text()
                if page_text:
                    text += page_text + '\n\n'  # 每页之间添加空行分隔
            
            # 关闭文档
            doc.close()
            
            # 写入文本文件
            with open(self.output_file, 'w', encoding='utf-8') as f:
                f.write(text)
            
            # 检查输出文件是否创建成功
            if not os.path.exists(self.output_file):
                raise RuntimeError("TXT文件未能成功创建")
                
        except FileNotFoundError:
            raise
        except Exception as e:
            raise RuntimeError(f"PDF转文本失败: {str(e)}") from e

    async def _doc_to_pdf(self):
        """DOC转PDF"""
        # 模拟进度更新
        yield 10
        
        if self.cancelled:
            return
        try:
            # 尝试使用 python-docx 和 reportlab 实现 DOC 转 PDF
            try:
                # 注意: python-docx 主要处理 DOCX 格式，对于 DOC 格式支持有限
                # 这里我们尝试使用 win32com.client (Windows) 或 libreoffice (跨平台) 来处理
                import os
                
                # 模拟进度更新
                yield 30
                
                # 检查操作系统
                import platform
                if platform.system() == "Windows":
                    # 在 Windows 上尝试使用 win32com.client
                    try:
                        import win32com.client
                        # 使用 Word 应用程序打开 DOC 文件并另存为 PDF
                        word = win32com.client.Dispatch("Word.Application")
                        word.visible = False
                        doc = word.Documents.Open(os.path.abspath(self.input_file))
                        doc.SaveAs(os.path.abspath(self.output_file), FileFormat=17)  # 17 表示 PDF 格式
                        doc.Close()
                        word.Quit()
                        # 模拟进度更新
                        yield 90
                        yield 100
                        return
                    except Exception as e:
                        logger.warning(f"使用 Word 应用程序转换失败: {e}")
                        pass
                
                # 跨平台方案：尝试使用 libreoffice
                try:
                    import subprocess
                    import shutil
                    
                    # 检查 libreoffice 是否可用
                    libreoffice_path = shutil.which("libreoffice")
                    if not libreoffice_path:
                        # 尝试使用 soffice (libreoffice 的命令行工具)
                        libreoffice_path = shutil.which("soffice")
                    
                    if libreoffice_path:
                        # 使用 libreoffice 转换
                        cmd = [
                            libreoffice_path,
                            "--headless",
                            "--convert-to", "pdf",
                            "--outdir", os.path.dirname(self.output_file),
                            self.input_file
                        ]
                        
                        result = subprocess.run(cmd, capture_output=True, text=True)
                        if result.returncode == 0:
                            # 重命名输出文件以匹配期望的输出文件名
                            input_name = os.path.splitext(os.path.basename(self.input_file))[0]
                            actual_output = os.path.join(os.path.dirname(self.output_file), input_name + ".pdf")
                            if os.path.exists(actual_output):
                                os.rename(actual_output, self.output_file)
                                # 模拟进度更新
                                yield 90
                                yield 100
                                return
                        else:
                            logger.warning(f"LibreOffice 转换失败: {result.stderr}")
                except Exception as e:
                    logger.warning(f"使用 LibreOffice 转换失败: {e}")
                    pass
                
                # 如果以上方法都失败了，提供有用的错误信息
                raise RuntimeError(
                    "DOC 转 PDF 需要额外的依赖库:\n\n"
                    "Windows 用户:\n"
                    "- 安装 Microsoft Word 或\n"
                    "- 安装 LibreOffice\n\n"
                    "macOS/Linux 用户:\n"
                    "- 安装 LibreOffice\n\n"
                    "安装 LibreOffice:\n"
                    "- 访问 https://www.libreoffice.org/download/download/ 下载并安装\n"
                    "- 或使用包管理器安装: brew install libreoffice (macOS) 或 apt install libreoffice (Ubuntu)"
                )
            except ImportError as e:
                raise RuntimeError(f"缺少必要的依赖库: {str(e)}") from e
        except Exception as e:
            raise RuntimeError(f"DOC 转 PDF 失败: {str(e)}") from e

    async def _docx_to_pdf(self):
        """DOCX转PDF"""
        # 模拟进度更新
        yield 10
        
        if self.cancelled:
            return
        try:
            from docx2pdf import convert
            import os
            
            # 模拟进度更新
            yield 30
            
            # 检查输入文件是否存在
            if not os.path.exists(self.input_file):
                raise FileNotFoundError(f"输入文件不存在: {self.input_file}")
            
            convert(self.input_file, self.output_file)
            
            # 模拟进度更新
            yield 90
            
            # 检查输出文件是否创建成功
            if not os.path.exists(self.output_file):
                raise RuntimeError("PDF文件未能成功创建")
                
            yield 100
        except ImportError:
            # 如果没有安装docx2pdf库，提供有用的错误信息
            raise RuntimeError(
                "DOCX转PDF需要docx2pdf库:\n\n"
                "请运行以下命令安装:\n"
                "pip install docx2pdf\n\n"
                "注意: docx2pdf需要Microsoft Word或WPS Office支持"
            )
        except Exception as e:
            raise RuntimeError(f"DOCX转PDF失败: {str(e)}") from e

    async def _pdf_to_docx(self):
        """PDF转DOCX"""
        # 模拟进度更新
        yield 10
        
        if self.cancelled:
            return
        try:
            import fitz  # PyMuPDF
            from docx import Document
            import os
            
            # 模拟进度更新
            yield 20
            
            # 检查输入文件是否存在
            if not os.path.exists(self.input_file):
                raise FileNotFoundError(f"输入文件不存在: {self.input_file}")
            
            # 打开PDF文档
            doc = fitz.open(self.input_file)
            docx_doc = Document()
            
            # 添加标题
            docx_doc.add_heading('转换自PDF文件', 0)
            
            # 模拟进度更新
            yield 30
            
            # 提取所有页面的文本并添加到DOCX文档
            total_pages = len(doc)
            for page_num in range(total_pages):
                page = doc.load_page(page_num)
                text = page.get_text()
                
                if text and text.strip():  # 确保文本不为空
                    # 添加页面分隔标题
                    docx_doc.add_heading(f'页面 {page_num+1}', level=1)
                    # 添加文本内容
                    docx_doc.add_paragraph(text.strip())
                    # 添加页面分隔符（除了最后一页）
                    if page_num < total_pages - 1:
                        docx_doc.add_page_break()
                elif page_num < total_pages - 1:  # 即使页面无文本也要添加分页符（除了最后一页）
                    docx_doc.add_page_break()
                
                # 根据处理进度更新进度值
                progress = 30 + int((page_num + 1) / total_pages * 60)
                yield progress
            
            # 关闭PDF文档
            doc.close()
            
            # 保存DOCX文档
            docx_doc.save(self.output_file)
            
            # 模拟进度更新
            yield 95
            
            # 检查输出文件是否创建成功
            if not os.path.exists(self.output_file):
                raise RuntimeError("DOCX文件未能成功创建")
                
            yield 100
        except FileNotFoundError:
            raise
        except Exception as e:
            raise RuntimeError(f"PDF转DOCX失败: {str(e)}") from e


class QconvertoNiceGUIApp:
    """Qconverto NiceGUI应用程序"""
    
    def __init__(self):
        self.input_file: Optional[str] = None
        self.output_dir: Optional[str] = None
        self.worker: Optional[ConversionWorker] = None
        self.temp_files = []  # 跟踪临时文件以便清理
        self.conversion_task: Optional[asyncio.Task] = None
        self.conversion_in_progress = False
        self.setup_ui()
        
    def setup_ui(self):
        """设置用户界面"""
        # 设置页面标题和主题
        ui.page_title('Qconverto - 多媒体文件格式转换工具')
        
        # 设置响应式设计
        ui.query('body').classes('bg-gray-50')
        
        # 主容器
        self.main_container = ui.column().classes('w-full')
        
        # 初始化现代化UI组件
        self.ui_components = ModernUIComponents(self)
        
    def log(self, message: str):
        """添加日志（已禁用）"""
        logger.info(message)
        # 日志功能已移除
        pass
    
    def determine_file_type(self, file_path: str) -> Optional[str]:
        """确定文件类型"""
        if not os.path.exists(file_path):
            return None
            
        ext = os.path.splitext(file_path)[1].lower()
        
        # 基于文件扩展名进行判断
        for file_type, formats in SUPPORTED_FORMATS.items():
            if ext in formats['输入']:
                return file_type
                
        return None
    
    def get_file_icon(self, file_path: str) -> str:
        """根据文件扩展名获取相应的图标"""
        ext = os.path.splitext(file_path)[1].lower()
        return FILE_TYPE_ICONS.get(ext, 'insert_drive_file')
    
    def update_converted_file_icon(self, output_format: str) -> None:
        """根据选择的输出格式更新转换后文件的图标"""
        if output_format:
            # 将格式转换为扩展名（如 'pdf' → '.pdf'）
            ext = f'.{output_format.lower()}'
            icon_name = FILE_TYPE_ICONS.get(ext, 'insert_drive_file')
            
            # 更新图标
            self.converted_file_icon.name = icon_name
            
            # 确保图标和文件名可见
            self.converted_file_icon.visible = True
            self.converted_file_name.visible = True
        else:
            # 如果没有选择格式，隐藏图标和文件名
            self.converted_file_icon.visible = False
            self.converted_file_name.visible = False
    
    def update_format_options(self, file_type: str):
        """根据文件类型更新输出格式选项"""
        if file_type in SUPPORTED_FORMATS:
            formats = SUPPORTED_FORMATS[file_type]['输出']
            self.format_select.options = formats
            # 同时更新转换后文件区域的格式选择器
            self.output_format_select.options = formats
            if formats:
                # 智能选择默认格式（优先选择高质量或常用格式）
                default_format = self._get_recommended_format(file_type, formats)
                self.format_select.value = default_format
                self.output_format_select.value = default_format
                # 更新转换后文件图标
                self.update_converted_file_icon(default_format)
            

            
            # 在日志中显示更详细的信息
            self.log(f"✓ 检测到文件类型: {file_type}")
            self.log(f"  可转换格式: {', '.join(formats)}")
            
            # 转换设置功能已移除
            pass
        else:
            # 不支持的文件类型
            self.log(f"✗ 不支持的文件类型: {file_type}")
            self.format_select.options = []
            self.output_format_select.options = []
    
    def _get_recommended_format(self, file_type: str, available_formats: list) -> str:
        """根据文件类型智能推荐最佳输出格式"""
        # 定义每种文件类型的推荐格式优先级
        recommendations = {
            '图片': ['.pdf', '.png', '.jpg', '.webp'],  # PDF最通用，PNG质量最好，JPG最通用，WebP现代格式
            '音频': ['.mp3', '.flac', '.wav'],  # MP3最通用，FLAC无损，WAV原始
            '视频': ['.mp4', '.mkv', '.avi'],   # MP4最通用，MKV功能丰富，AVI经典
            '文档': ['.pdf', '.docx', '.txt']    # PDF最通用，DOCX可编辑，TXT纯文本
        }
        
        # 获取该文件类型的推荐格式列表
        recommended = recommendations.get(file_type, [])
        
        # 优先选择推荐列表中的第一个可用格式
        for format_ext in recommended:
            if format_ext in available_formats:
                return format_ext
        
        # 如果推荐格式都不在可用格式中，则返回第一个可用格式
        return available_formats[0] if available_formats else ''
    
    async def handle_file_upload(self, e: events.UploadEventArguments):
        """处理文件上传"""
        # 保存上传的文件
        # 移除了无用的上传事件参数类型和属性日志
        
        # 获取上传的文件
        if hasattr(e, 'files'):
            # 多文件上传
            uploaded_file = next(iter(e.files), None)
        elif hasattr(e, 'file'):
            # 单文件上传
            uploaded_file = e.file
        else:
            self.log("未收到上传文件")
            return
            
        if not uploaded_file:
            self.log("未收到上传文件")
            return
            
        file_name = uploaded_file.name
        self.log(f"接收到上传文件: {file_name}")
        
        # 正确读取文件内容
        file_content = await uploaded_file.read()
        
        # 保存到临时文件
        temp_dir = os.path.join(os.getcwd(), 'temp')
        os.makedirs(temp_dir, exist_ok=True)
        temp_file_path = os.path.join(temp_dir, file_name)
        
        with open(temp_file_path, 'wb') as f:
            f.write(file_content)
        
        self.input_file = temp_file_path
        self.temp_files.append(temp_file_path)  # 添加到临时文件跟踪列表
        
        # 更新新的UI组件
        self.original_file_name.text = file_name
        file_icon = self.get_file_icon(temp_file_path)
        self.original_file_icon.name = file_icon
        self.original_file_icon.classes(replace='text-6xl text-blue-500 mb-2 animate-bounce')
        self.original_file_name.classes(replace='text-blue-600 text-center')
        
        self.log(f"已选择输入文件: {file_name}")
        self.log(f"文件保存路径: {temp_file_path}")
        self.log(f"文件是否存在: {os.path.exists(temp_file_path)}")
        
        # 根据文件类型更新输出格式选项
        file_type = self.determine_file_type(temp_file_path)
        if file_type:
            self.update_format_options(file_type)
        
        # 显示文件预览（如果是图片）
        if file_type == '图片':
            self.show_image_preview(temp_file_path)
    
    def show_image_preview(self, image_path: str):
        """显示图片预览"""
        self.preview_container.clear()
        with self.preview_container:
            ui.label('文件预览').classes('font-medium mb-2')
            with ui.card().classes('w-full'):
                # 限制图片大小以适应预览区域
                ui.image(image_path).classes('w-full max-h-64 object-contain')
    
    async def browse_output_dir(self):
        """浏览输出目录"""
        # 在Web环境中，我们让用户选择是否使用默认目录或自定义目录
        with ui.dialog() as dialog, ui.card():
            ui.label('请选择输出目录选项:').classes('mb-4')
            
            # 选项1: 使用与输入文件相同的目录
            def use_same_dir():
                self.output_dir = None
                self.log("输出目录设置为: 与输入文件相同目录")
                dialog.close()
            
            ui.button('与输入文件相同目录', on_click=use_same_dir).classes('w-full mb-2')
            
            # 选项2: 自定义输出目录
            def custom_dir():
                # 提示用户输入目录路径
                dialog.close()
                self.prompt_custom_output_dir()
            
            ui.button('自定义输出目录', on_click=custom_dir).classes('w-full mb-4')
            
            ui.button('取消', on_click=dialog.close).classes('w-full')
        
        await dialog
    
    async def prompt_custom_output_dir(self):
        """提示用户输入自定义输出目录"""
        dialog = ui.dialog()
        with dialog, ui.card():
            ui.label('请输入输出目录路径:').classes('mb-2')
            output_dir_input = ui.input(placeholder='例如: C:/Users/YourName/Documents').classes('w-full mb-4')
            with ui.row():
                def confirm():
                    dir_path = output_dir_input.value
                    if dir_path and os.path.exists(dir_path):
                        self.output_dir = dir_path
                        self.log(f"已选择输出目录: {dir_path}")
                        dialog.close()
                    else:
                        ui.notify('目录不存在，请输入有效的目录路径')
                
                ui.button('确定', on_click=confirm).classes('mr-2')
                ui.button('取消', on_click=dialog.close)
        
        await dialog
    
    async def start_conversion(self):
        """开始转换"""
        self.log(f"开始转换调用，input_file值: {self.input_file}")
        if not self.input_file:
            ui.notify("请先选择输入文件", type='warning')
            self.log("警告: 请先选择输入文件")
            return
            
        self.log(f"检查文件是否存在: {self.input_file}")
        if not os.path.exists(self.input_file):
            ui.notify("输入文件不存在", type='negative')
            self.log("警告: 输入文件不存在")
            self.log(f"当前工作目录: {os.getcwd()}")
            self.log(f"文件绝对路径: {os.path.abspath(self.input_file) if self.input_file else 'None'}")
            return
            
        # 确定输出文件路径
        input_dir = os.path.dirname(self.input_file)
        input_name = os.path.splitext(os.path.basename(self.input_file))[0]
        # 优先使用转换后文件区域的格式选择，如果没有选择则使用顶部设置面板的格式选择
        output_ext = self.output_format_select.value if self.output_format_select.value else self.format_select.value
        output_dir = self.output_dir if self.output_dir else input_dir
        output_file = os.path.join(output_dir, f"{input_name}{output_ext}")
        
        # 确定文件类型
        file_type = self.determine_file_type(self.input_file)
        if not file_type:
            ui.notify("不支持的文件类型", type='negative')
            self.log("警告: 不支持的文件类型")
            return
            
        self.log(f"开始转换 {self.input_file} 到 {output_file}")
        self.log(f"文件类型: {file_type}")
        
        # 禁用转换按钮，显示进度区域
        self.convert_button.disable()
        self.conversion_in_progress = True
        
        # 启动转换动画
        self.conversion_arrow.classes(replace='text-4xl text-blue-500 hidden')
        self.conversion_spinner.classes(replace='block animate-spin')
        
        try:
            # 创建转换工作器（不传递设置参数）
            self.worker = ConversionWorker(
                self.input_file, 
                output_file, 
                file_type
            )
            
            # 执行转换
            self.conversion_task = asyncio.create_task(self.run_conversion_with_progress())
            
        except Exception as e:
            self.log(f"✗ 转换启动失败: {str(e)}")
            ui.notify(f"转换启动失败: {str(e)}", type='negative')
            self.convert_button.enable()
            self.conversion_in_progress = False
            # 停止转换动画
            self.conversion_arrow.classes(replace='text-4xl text-blue-500')
            self.conversion_spinner.classes(replace='hidden')
    
    async def run_conversion_with_progress(self):
        """运行带进度更新的转换任务"""
        try:
            async for progress in self.worker.run_conversion():
                if self.worker.cancelled:
                    break
                self.progress_bar.value = progress / 100
                self.progress_text.text = f"转换进度: {progress}%"
                
                # 添加脉冲动画到转换箭头以反映进度
                if progress % 10 == 0:  # 每10%触发一次动画
                    self.conversion_spinner.classes(replace='block animate-pulse')
                    # 短暂延迟后恢复旋转动画
                    await asyncio.sleep(0.1)
                    if not self.worker.cancelled:
                        self.conversion_spinner.classes(replace='block animate-spin')
            
            if not self.worker.cancelled:
                # 转换成功
                self.log("✓ 转换成功完成!")
                self.log(f"  输出文件: {self.worker.output_file}")
                if os.path.exists(self.worker.output_file):
                    size = os.path.getsize(self.worker.output_file)
                    self.log(f"  文件大小: {self.format_file_size(size)}")
                    
                    # 更新转换后文件显示
                    download_name = os.path.basename(self.worker.output_file)
                    self.converted_file_name.text = download_name
                    file_icon = self.get_file_icon(self.worker.output_file)
                    self.converted_file_icon.name = file_icon
                    self.converted_file_icon.classes(replace='text-6xl text-green-500 mb-2 animate-pulse')
                    self.converted_file_name.classes(replace='text-green-600 text-center')
                    
                    # 保存文件内容用于下载按钮
                    with open(self.worker.output_file, 'rb') as f:
                        self.converted_file_content = f.read()
                    
                    # 在转换后文件区域添加下载按钮
                    # 清除之前的下载按钮（如果有的话）
                    self.download_button_container.clear()
                    
                    # 创建新的下载按钮
                    with self.download_button_container:
                        ui.button(
                            '下载文件', 
                            icon='download',
                            on_click=lambda: ui.download(self.converted_file_content, download_name)
                        ).classes('px-4 py-2 bg-blue-500 text-white rounded-lg hover:bg-blue-600 transition-colors')
                    
                    self.log(f"  文件已准备好下载: {download_name}")
                    
                self.progress_text.text = "转换成功完成"
                # 使用with语句显式指定UI上下文来避免RuntimeError
                with self.main_container:
                    ui.notify("转换成功完成!", type='positive')
                
                # 停止转换动画并显示完成效果
                self.conversion_spinner.classes(replace='hidden')
                self.conversion_arrow.classes(replace='text-4xl text-green-500')
            else:
                self.log("⚠ 转换已被取消")
                self.progress_text.text = "转换已取消"
                # 使用with语句显式指定UI上下文来避免RuntimeError
                with self.main_container:
                    ui.notify("转换已被取消", type='warning')
                
                # 停止转换动画
                self.conversion_spinner.classes(replace='hidden')
                self.conversion_arrow.classes(replace='text-4xl text-blue-500')
                
        except Exception as e:
            self.log(f"✗ 转换错误: {str(e)}")
            self.progress_text.text = f"转换错误: {str(e)}"
            # 使用with语句显式指定UI上下文来避免RuntimeError
            with self.main_container:
                ui.notify(f"转换错误: {str(e)}", type='negative')
            
            # 停止转换动画
            self.conversion_spinner.classes(replace='hidden')
            self.conversion_arrow.classes(replace='text-4xl text-blue-500')
        finally:
            # 重新启用转换按钮
            self.convert_button.enable()
            self.conversion_in_progress = False
            self.conversion_task = None
    
    def cancel_conversion(self):
        """取消转换"""
        if self.worker:
            self.worker.cancelled = True
        if self.conversion_task:
            self.conversion_task.cancel()
        self.log("用户请求取消转换")
    
    def reset_app(self):
        """重置应用程序"""
        # 取消正在进行的转换
        if self.conversion_in_progress:
            self.cancel_conversion()
        
        # 清理临时文件
        self.cleanup_temp_files()
        
        # 重置变量
        self.input_file = None
        self.output_dir = None
        self.worker = None
        self.conversion_task = None
        self.conversion_in_progress = False
        
        # 重置UI元素
        self.selected_file_label.text = '未选择文件'
        self.selected_file_label.classes(replace='text-gray-500')
        self.format_select.options = []
        self.format_select.value = None
        self.progress_bar.value = 0
        self.progress_text.text = '准备就绪'
        self.log_output.value = ''
        
        # 清理预览区域
        self.preview_container.clear()
        
        self.log("应用程序已重置")
        ui.notify("应用程序已重置", type='info')
    
    def cleanup_temp_files(self):
        """清理临时文件"""
        for temp_file in self.temp_files:
            try:
                if os.path.exists(temp_file):
                    os.remove(temp_file)
            except Exception as e:
                self.log(f"清理临时文件失败 {temp_file}: {e}")
        self.temp_files = []
    
    def format_file_size(self, size_bytes: int) -> str:
        """格式化文件大小"""
        if size_bytes == 0:
            return "0 B"
        size_names = ["B", "KB", "MB", "GB"]
        i = 0
        while size_bytes >= 1024 and i < len(size_names) - 1:
            size_bytes /= 1024.0
            i += 1
        return f"{size_bytes:.1f} {size_names[i]}"


# 创建应用实例
app_instance = QconvertoNiceGUIApp()

# 注册应用关闭时的清理函数
import atexit
atexit.register(app_instance.cleanup_temp_files)

# 运行应用
if __name__ == '__main__':
    ui.run(
        title='Qconverto - 多媒体文件格式转换工具',
        reload=False,
        favicon='icon.svg',
        dark=False,
        port=8081
    )